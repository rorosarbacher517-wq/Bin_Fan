from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "docs" / "GridWeather-Agent_秋招面试项目详解手册.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(31, 41, 55)
MUTED = RGBColor(90, 98, 110)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[float]) -> None:
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
            set_cell_margins(row.cells[idx])
            row.cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")


def add_para(doc: Document, text: str = "", style: str | None = None, bold_prefix: str | None = None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbers(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_kv_table(doc: Document, rows: list[tuple[str, str]], widths=(1.65, 4.85)) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_width(table, list(widths))
    hdr = table.rows[0].cells
    hdr[0].text = "维度"
    hdr[1].text = "内容"
    for c in hdr:
        set_cell_shading(c, LIGHT_BLUE)
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
    for k, v in rows:
        cells = table.add_row().cells
        cells[0].text = k
        cells[1].text = v
    doc.add_paragraph()


def add_matrix(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_width(table, widths)
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        set_cell_shading(table.rows[0].cells[i], LIGHT_BLUE)
        for p in table.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    doc.add_paragraph()


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [6.5])
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, CALLOUT)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = DARK_BLUE
    cell.add_paragraph(body)
    doc.add_paragraph()


def set_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        st = doc.styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.line_spacing = 1.25

    for name in ["List Bullet", "List Number"]:
        st = doc.styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(11)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.line_spacing = 1.25


def add_title_page(doc: Document) -> None:
    section = doc.sections[0]
    section.header.paragraphs[0].text = "GridWeather-Agent 面试项目手册"
    section.header.paragraphs[0].style = doc.styles["Header"]
    footer = section.footer.paragraphs[0]
    footer.text = "Generated for AI Lab / 电网 / 大厂算法 / Agent / 多模态 / AI Infra 秋招准备"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("GridWeather-Agent")
    r.font.size = Pt(26)
    r.bold = True
    r.font.color.rgb = DARK_BLUE

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run("多模态电网微气象风险预测与可解释 Agent 系统：工程实现、科学问题与秋招面试手册")
    r.font.size = Pt(14)
    r.font.color.rgb = MUTED

    add_kv_table(
        doc,
        [
            ("项目定位", "AI for Earth and Energy：真实气象再分析 + 真实遥感/DEM + 非敏感线路 + 物理弱标签 + 可解释风险报告。"),
            ("研究区", "贵州毕节-六盘水山地走廊，bbox = [27.7N, 104.3E, 26.1N, 106.1E]。"),
            ("真实数据", "Copernicus ERA5-Land；Google Earth Engine Sentinel-2 SR Harmonized；NASA/NASADEM。"),
            ("当前结果", "70,560 条训练样本；3,360 条最新预测；accuracy = 0.908；macro-F1 = 0.877；risk score MAE = 1.030。"),
            ("使用方式", "既可讲科学问题，也可讲数据工程、模型算法、Agent 工具链和 AI Infra。"),
        ],
    )
    add_callout(
        doc,
        "一句话项目介绍",
        "我构建了一个面向贵州复杂山地输电线路的多模态微气象风险预测系统，自动下载 ERA5-Land 气象数据，利用 GEE 提取 DEM 和 Sentinel-2 静态地表特征，结合非敏感合成线路和物理弱标签训练风险模型，并输出杆塔级风险解释报告。",
    )
    doc.add_page_break()


def add_toc(doc: Document) -> None:
    doc.add_heading("阅读路线", level=1)
    add_numbers(
        doc,
        [
            "先掌握项目叙事：为什么这个问题比普通天气预测更适合行业落地。",
            "再掌握工程链路：从 CDS/GEE 到训练表、模型、报告，每一步能解释输入输出。",
            "再掌握算法与数学：弱标签、风险分级、梯度提升、评价指标、时空对齐。",
            "最后按岗位准备问题：AI Lab、电网、大厂算法、Agent、多模态、AI Infra 和 HR 面。",
        ],
    )


def add_project_story(doc: Document) -> None:
    doc.add_heading("一、项目定位与原项目互补关系", level=1)
    add_para(
        doc,
        "这个项目不是替代原有 HLS 碳通量 / footprint physics 项目，而是补齐工程落地与行业场景。原项目证明你有科学研究深度；GridWeather-Agent 证明你能把遥感、气象、地理空间数据和 AI 工程做成可运行系统。",
    )
    add_matrix(
        doc,
        ["维度", "HLS 碳通量项目", "GridWeather-Agent 项目"],
        [
            ["核心科学问题", "遥感像元、站点通量和 footprint 尺度如何匹配。", "格点气象如何下推到输电线路杆塔级灾害风险。"],
            ["数据主线", "HLS、通量站、生态系统变量、footprint 权重。", "ERA5-Land、Sentinel-2、NASADEM、线路/杆塔点。"],
            ["模型主线", "生态遥感回归、空间机制解释、统计评估。", "多源时空融合、物理弱监督、风险分类/回归、Agent 解释。"],
            ["工程主线", "科研数据处理和论文图表。", "自动下载、GEE 提取、训练、推理、报告、可部署接口。"],
            ["适配岗位", "AI for Science、遥感生态、科研算法。", "电网能源、行业算法、多模态、Agent、AI Infra。"],
        ],
        [1.25, 2.55, 2.7],
    )
    add_callout(
        doc,
        "个人标签",
        "你可以把两条线合并为：AI for Earth and Energy，擅长将遥感、气象、地理空间数据与物理机制结合，构建可解释、可复现、可工程化部署的时空智能系统。",
    )


def add_industry_need(doc: Document) -> None:
    doc.add_heading("一、为什么原 MVP 还浅：真实机构需求反推", level=1)
    add_para(
        doc,
        "如果只做“ERA5 + DEM + Sentinel + 分类器预测覆冰风险”，面试官会觉得它是一个可运行 demo，但距离 AI Lab、电网和大厂核心需求还有距离。真实需求不是单点分类，而是极端天气下电网运行韧性：天气会不会造成灾害，灾害会不会影响线路能力，系统应该如何解释和建议处置。",
    )
    add_matrix(
        doc,
        ["外部趋势", "公开依据", "对项目的设计要求"],
        [
            ["电网成为能源转型瓶颈", "IEA 电网报告强调电气化、热泵、电动车和新能源接入增加电网压力，电网规划与管理方式也需要升级。", "项目不能只预测天气，要连接到线路能力、运维优先级和可靠供电。"],
            ["AI 与能源深度耦合", "IEA Energy and AI 报告指出 AI 既会增加电力需求，也可能改变能源行业运行方式。", "项目要展示 AI for energy optimization，而不只是 AI 使用者。"],
            ["线路评级从静态转向天气感知", "FERC Order 881 等政策背景强调输电线路额定值与环境条件相关。", "加入 DLR/RTTR 思想：温度、风、线路夹角影响动态容量余量。"],
            ["天气大模型仍有极端事件短板", "GenCast/AIFS 证明天气 AI 快速发展，但 record-breaking extremes 研究指出 AI 天气模型对前所未见极端事件仍有风险。", "项目定位为下游高影响天气诊断和决策支持，主动做不确定性和物理约束。"],
            ["杆塔级风险需要遥感/DEM", "ERA5-Land 0.1 度/小时，Sentinel-2 和 NASADEM 提供更细下垫面和地形信息。", "用 ERA5 做背景场，用 DEM/Sentinel/线路做局地修正。"],
        ],
        [1.25, 2.45, 2.8],
    )
    add_callout(
        doc,
        "升级后的项目定义",
        "GridWeather-Agent 不再只是覆冰风险分类器，而是 Weather-to-Grid Resilience Agent：融合真实气象、真实遥感/DEM、线路几何和动态线路容量代理，输出杆塔级灾害风险、容量余量、证据链和运维动作建议。",
    )


def add_architecture(doc: Document) -> None:
    doc.add_heading("三、端到端工程架构", level=1)
    add_para(doc, "当前工程位于 work/GridWeatherAgent，已经形成从数据到报告的闭环。核心层次如下：")
    add_matrix(
        doc,
        ["层次", "模块/文件", "输入", "输出", "面试考点"],
        [
            ["数据下载", "download_era5_land.py", "CDS API、project.yaml", "ERA5-Land NetCDF/ZIP", "API 鉴权、请求拆分、队列任务、重试、数据许可。"],
            ["格式转换", "prepare_era5_weather.py / era5_to_weather.py", "ERA5 NetCDF/ZIP", "weather_hourly.csv", "NetCDF、xarray、单位转换、风速风向、相对湿度。"],
            ["遥感/地形", "export_gee_static_features.py", "塔点 CSV、GEE project", "tower_static_features_gee.csv", "GEE、Sentinel-2、NASADEM、buffer reduceRegions、代理与 OAuth。"],
            ["线路生成", "synthetic.py", "研究区 bbox、随机种子", "非敏感杆塔点", "涉密规避、空间采样、可复现。"],
            ["特征融合", "build_dataset.py", "weather + towers + static", "training_table.csv", "时空 join、最近格点匹配、静态/动态特征融合。"],
            ["动态线路能力", "dlr.py", "温度、风速、线路风夹角", "dlr_ampacity_a、dlr_margin_pct", "DLR/RTTR、导线热平衡、天气到运行能力。"],
            ["弱标签", "physics_labels.py", "温度、湿度、风、降水、海拔、坡度", "risk_score/risk_level", "物理机理、弱监督、风险分箱。"],
            ["模型训练", "train.py", "training_table.csv", "risk_model.joblib + metrics", "分类、回归、时间切分、macro-F1、过拟合。"],
            ["推理报告", "predict.py / report.py / explain.py", "最新 24h 样本", "latest_predictions.csv + HTML", "批量推理、可解释规则、业务报告。"],
        ],
        [0.9, 1.35, 1.15, 1.15, 1.95],
    )


def add_data_steps(doc: Document) -> None:
    doc.add_heading("四、从下载到出结果的详细步骤", level=1)
    steps = [
        ("步骤 1：确定研究区和时间窗", "选择贵州毕节-六盘水山地走廊，时间窗为 2023-01-01 至 2023-01-21。选择原因是贵州山地高海拔、坡度和迎风坡差异明显，冬季凝冻/覆冰风险具有业务意义。"),
        ("步骤 2：下载 ERA5-Land", "使用 CDS API 请求 reanalysis-era5-land，变量包括 2m temperature、2m dewpoint temperature、10m u/v wind、total precipitation、surface pressure。脚本按月份生成请求，CDS 返回后保存到 data/demo/raw/era5_land。"),
        ("步骤 3：处理 CDS 返回格式", "CDS 新接口可能返回 ZIP 容器，即使文件名是 .nc。转换脚本会自动检测 zipfile，解压 data_0.nc，再由 xarray 读取。"),
        ("步骤 4：气象变量单位转换", "温度和露点从 K 转为摄氏度；风速由 sqrt(u10^2 + v10^2) 得到；风向由 arctan2(-u, -v) 转为 0-360 度；降水从 m 转为 mm；气压从 Pa 转为 hPa。"),
        ("步骤 5：GEE 提取地形和遥感静态特征", "使用 NASADEM 提取 elevation、slope、aspect；使用 Sentinel-2 SR Harmonized 计算 NDVI、NDWI、NDBI；对每个杆塔 500 m buffer 做均值 reduceRegions。"),
        ("步骤 6：构造非敏感杆塔线路", "在研究区内按固定随机种子生成 4 条合成线路、每条 35 个杆塔。这样避免真实电网坐标涉密，同时保留线路走向、海拔、地形暴露等建模要素。"),
        ("步骤 7：时空对齐", "将每个杆塔匹配到最近 ERA5-Land 格点，动态气象按小时重复到杆塔级，静态 DEM/Sentinel 特征按 tower_id 合并。"),
        ("步骤 8：计算动态线路容量余量", "用环境温度、风速和风向-线路夹角估算导线冷却能力，得到 dlr_ampacity_a、dlr_margin_pct 和 thermal_stress_index。它不是认证级 IEEE 738 实现，而是面试项目中的 DLR/RTTR 物理代理。"),
        ("步骤 9：生成物理弱标签", "用低温窗口、高湿、降水、风速、海拔、坡度构造 risk_score，再分箱为 0/1/2/3 风险等级。该标签不是人工真值，而是可复现的物理先验弱监督。"),
        ("步骤 10：时间切分训练模型", "按时间留出最后 5 天做测试，避免随机切分造成时间泄漏。模型同时训练风险等级分类器和风险分数回归器。"),
        ("步骤 11：推理和解释", "对最新 24 小时所有杆塔批量推理，输出 risk_score、risk_level、DLR 余量，并生成触发因素和建议动作。"),
    ]
    for title, body in steps:
        doc.add_heading(title, level=2)
        add_para(doc, body)


def add_algorithms(doc: Document) -> None:
    doc.add_heading("五、模型、算法与数学细节", level=1)
    doc.add_heading("4.1 当前工程实际使用的模型", level=2)
    add_matrix(
        doc,
        ["模型/算法", "用途", "为什么用它", "面试追问"],
        [
            ["HistGradientBoostingClassifier", "预测 risk_level 0/1/2/3", "适合表格特征、训练快、可处理非线性阈值关系。", "和 XGBoost/LightGBM 区别？为什么不用深度模型？"],
            ["HistGradientBoostingRegressor", "预测连续 risk_score", "给报告提供更细粒度分数，避免只有离散等级。", "分类和回归为什么同时做？"],
            ["ColumnTransformer", "数值特征标准化、line_id one-hot", "工程化 pipeline，训练和推理保持同样预处理。", "如何防止训练/推理特征不一致？"],
            ["物理弱标签", "构造可复现风险监督信号", "解决真实电网故障标签涉密/稀缺问题。", "弱标签会不会带偏模型？如何验证？"],
            ["DLR 物理代理", "估算天气驱动的容量余量", "把 hazard prediction 连接到 grid operation。", "和 IEEE 738/CIGRE 正式热平衡有什么差异？"],
            ["规则解释 Agent", "输出触发原因和建议动作", "业务侧需要知道为什么高风险、该做什么。", "规则解释和模型解释 SHAP 有什么区别？"],
        ],
        [1.4, 1.15, 1.85, 2.1],
    )
    doc.add_heading("4.2 风险分数数学形式", level=2)
    add_para(
        doc,
        "当前 risk_score 是可解释的物理启发组合，不是神秘黑箱。其思想是将覆冰风险拆成温度窗口、湿度、降水、风速和地形暴露五个因子：",
    )
    add_bullets(
        doc,
        [
            "温度窗口：覆冰高风险通常出现在 -8 到 1.5 摄氏度附近，过暖不结冰，过冷且无液态水时增长受限。",
            "湿度因子：相对湿度越高，云雾/过冷水滴条件越可能满足。",
            "降水因子：冻雨、湿雪、云雾沉积都需要水分来源。",
            "风速因子：适中风速有利于水滴碰撞导线/杆塔，过小输送不足，过强可能改变沉积形态。",
            "地形因子：海拔、坡度、山脊/迎风坡会改变温湿风场。",
        ],
    )
    add_para(
        doc,
        "可以在面试中写成：score = 100 * (0.45*T_window + 0.20*moisture + 0.15*precip + 0.10*wind + 0.10*terrain)，并对温度过高、湿度过低或极冷少水情形做抑制。重点不是公式绝对真实，而是将领域先验显式编码为弱监督信号。",
    )
    doc.add_heading("4.3 评价指标", level=2)
    add_matrix(
        doc,
        ["指标", "当前值", "意义", "面试解释"],
        [
            ["Accuracy", "0.908", "总体等级预测准确率。", "类别不均衡时不能只看 accuracy。"],
            ["Macro-F1", "0.877", "四个风险等级的平均 F1。", "更关注高风险小类，不让低风险样本淹没指标。"],
            ["Risk score MAE", "1.030", "连续风险分数平均绝对误差。", "报告展示时连续分数比离散等级更细。"],
            ["Level 3 recall", "0.56", "最高风险等级召回。", "这是后续重点优化项，预警系统宁可多报也不能漏报。"],
            ["Capacity watch", "476 / 3,360", "最新 24h 中 DLR 余量小于 10% 的记录。", "把气象风险转成运行约束信号。"],
        ],
        [1.25, 0.75, 1.8, 2.7],
    )


def add_domain_knowledge(doc: Document) -> None:
    doc.add_heading("六、遥感、气象与电网业务知识", level=1)
    doc.add_heading("5.1 遥感/GEE 考点", level=2)
    add_bullets(
        doc,
        [
            "Sentinel-2 SR Harmonized 是地表反射率产品，适合计算 NDVI、NDWI、NDBI 等指数。",
            "NDVI = (NIR - Red) / (NIR + Red)，反映植被覆盖。植被和下垫面会影响近地层湿度、辐射和地表粗糙度。",
            "NDWI = (Green - NIR) / (Green + NIR)，常用于水体/湿润状况；本项目中更像静态水分环境 proxy。",
            "NDBI = (SWIR - NIR) / (SWIR + NIR)，通常反映建设用地/裸地倾向。",
            "NASADEM elevation、slope、aspect 反映复杂地形。坡向和海拔会影响局地温度、风速衰减和迎风凝结。",
            "GEE reduceRegions 的本质是将 raster 在每个 tower buffer 内聚合为表格特征。",
        ],
    )
    doc.add_heading("5.2 气象/覆冰考点", level=2)
    add_bullets(
        doc,
        [
            "ERA5-Land 是再分析数据，不是纯观测，也不是业务预报；它融合模型和观测，适合历史样本构建。",
            "2m temperature 和 dewpoint 可以估算相对湿度；露点越接近温度，空气越接近饱和。",
            "u10/v10 是东西向/南北向风分量，风速由向量模长计算，风向需注意气象方向定义。",
            "total precipitation 在 ERA5-Land 中是累积量/小时输出语义，实际使用时要确认时间步长和单位。",
            "输电线路覆冰风险受温度、湿度、降水、风速、海拔、地形、线路走向共同影响。",
            "贵州凝冻场景常与高海拔、低温高湿、山地迎风坡、冻雨/雾凇等因素有关。",
        ],
    )
    doc.add_heading("5.3 电网业务考点", level=2)
    add_bullets(
        doc,
        [
            "真实输电线路坐标、故障记录通常涉密，因此公开简历项目应使用合成线路或公开 OSM 数据。",
            "业务输出不能只有模型分数，还要有风险等级、触发原因、空间位置和巡检建议。",
            "预警系统评价不能只看平均准确率，还要看高风险召回、提前量、误报成本和漏报成本。",
            "杆塔级风险比格点级天气更贴近电网运维，因为风险沿线路高度不均一。",
            "动态线路评级 DLR/实时热额定 RTTR 的核心是导线允许电流随环境温度、风速、太阳辐射和导线温度上限变化。",
            "风险预警需要区分 hazard 和 impact：有覆冰风险不一定影响当前潮流，容量余量低时风险才更接近调度/运维决策。",
        ],
    )


def add_infra_agent(doc: Document) -> None:
    doc.add_heading("七、Agent、多模态与 AI Infra 讲法", level=1)
    add_matrix(
        doc,
        ["方向", "本项目对应实现", "可继续增强", "面试亮点"],
        [
            ["Agent", "规则解释模块根据预测、DLR 余量和触发因子生成动作建议。", "加入工具调用：查天气、查塔点、查负荷、生成报告、检索应急规则。", "Agent 不直接胡说，而是调用确定性工具和模型输出。"],
            ["多模态", "动态气象 + 静态遥感/DEM + 线路空间属性。", "加入遥感预训练模型、图神经网络线路拓扑、时序 Transformer。", "多模态不是拼接图片和文字，而是融合不同时空尺度的数据。"],
            ["AI Infra", "脚本化数据下载、转换、训练、推理、报告。", "FastAPI、Docker、Redis 缓存、PostGIS、ONNX、日志监控。", "能讲从 notebook 到服务化系统的工程演进。"],
            ["算法岗", "表格模型、弱监督、时间切分、指标评估。", "LightGBM/XGBoost、TFT/PatchTST、SHAP、校准曲线。", "先用强 baseline，再上复杂模型。"],
        ],
        [0.9, 1.8, 1.8, 2.0],
    )
    add_callout(
        doc,
        "Agent 面试核心口径",
        "这个项目中的 Agent 不是聊天包装，而是一个工具调用层：它读取模型预测、气象变量、地形遥感特征和规则库，生成有证据的风险解释。这样可以降低幻觉，并让业务人员看到可追溯的触发因子。",
    )


def add_production_engineering(doc: Document) -> None:
    doc.add_heading("八、工程思维与上线意识：从 Demo 到企业落地", level=1)
    add_para(
        doc,
        "企业面试不会只问模型分数。真正上线时，面试官会追问并发、接口延迟、token 成本、缓存、异常兜底、数据隐私、限流降级、检索压测和 Agent 上下文管理。当前工程新增了 infra/retrieval/memory/load_test 模块，目的就是把这些问题变成可运行代码。",
    )
    add_matrix(
        doc,
        ["工程问题", "本项目实现", "上线取舍", "面试官追问"],
        [
            ["接口延迟", "simulate_api_load.py 统计 p50/p95。", "模型常驻内存；静态特征预计算；批量推理。", "如何把 p95 从秒级降到毫秒/百毫秒级？"],
            ["缓存优化", "TTLLRUCache：TTL + LRU，O(1) get/set。", "天气/塔点/检索结果可缓存，预测结果按时间窗缓存。", "缓存 key 怎么设计？如何处理脏数据？"],
            ["限流降级", "TokenBucketRateLimiter。", "高峰时先拒绝低优先级请求；返回上次结果或低精度模型。", "令牌桶和漏桶区别？"],
            ["异常兜底", "retry 和 fallback wrapper。", "GEE/CDS/API 失败时用最近数据、缓存数据、规则模型兜底。", "哪些异常能重试，哪些不能？"],
            ["token 成本", "BudgetedConversationMemory 估算 token 并裁剪。", "LLM 只负责解释和报告，不直接做批量数值预测。", "如何控制长报告成本？"],
            ["上下文溢出", "summary memory + recent turns。", "长期事实摘要化，近轮对话保留，原始证据按需检索。", "什么时候写入记忆，什么时候删除？"],
            ["检索切片", "fixed_window_chunks / paragraph_chunks / BM25Retriever。", "规程类文档按章节/段落，日志类按时间窗口，表格按行/实体。", "chunk 太大/太小分别有什么问题？"],
            ["检索压测", "BM25 和 load test 可本地验证。", "评估 recall@k、MRR、延迟、空结果率、幻觉率。", "没有标注集怎么评估 RAG？"],
            ["数据隐私", ".gitignore 排除 data/artifacts/key。", "真实电网坐标和故障记录不进公开仓库；脱敏或合成线路。", "如何做权限隔离和审计？"],
        ],
        [1.05, 1.55, 1.95, 1.95],
    )
    doc.add_heading("8.6 本轮增强项落地状态", level=2)
    add_matrix(
        doc,
        ["增强项", "已加入的工程文件", "当前状态", "面试讲法"],
        [
            ["系统消融", "scripts/experiments/run_ablation.py", "已跑出 ablation_results.csv。最佳组为 weather+DEM+Sentinel+line+IEEE738，macro-F1=0.886。", "用实验说明哪些模态真的有效。"],
            ["解释图", "scripts/experiments/permutation_importance.py", "已生成 permutation_importance.png/csv。", "SHAP 未安装，用 permutation importance 作为无依赖解释 baseline。"],
            ["模型升级", "patchtst_lite.py / train_patchtst_lite.py", "代码已实现；当前本机 torch DLL 异常，脚本优雅 skipped。", "诚实说明深度模型接口已准备，环境满足即可跑。"],
            ["模型工厂", "models/model_zoo.py", "LightGBM/XGBoost 可选导入；未安装时 fallback 到 HistGB。", "展示工程兼容性而不是强依赖环境。"],
            ["线路拓扑", "graph_smooth_predictions.py", "已实现线路 path graph 风险平滑后处理。", "可解释的 GNN 前置版本，后续替换为图神经网络消息传递。"],
            ["IEEE738/CIGRE DLR", "features/ieee738.py", "加入空气密度、太阳增益、对流/辐射冷却、ampacity、margin。", "明确是 IEEE738-like study model，不冒充认证级。"],
            ["服务化", "service/api.py、deploy/Dockerfile、docker-compose.yml", "FastAPI/Redis/PostGIS 为可选部署依赖，本机未安装 FastAPI。", "展示服务接口、健康检查、缓存、限流设计。"],
            ["Agent/RAG", "agent/toolchain.py、guidelines corpus", "已接入运维规程 hybrid retrieval 和 evidence packet。", "Agent 是工具链，不是聊天壳。"],
            ["压测增强", "simulate_api_load.py、circuit_breaker.py、singleflight.py", "已有限流、缓存、熔断、防击穿、本地压测。", "能讲高并发保护和降级策略。"],
            ["检索增强", "hybrid.py、evaluate_chunking.py", "BM25 + TF-IDF hybrid + rerank + chunking eval。", "能比较 chunk 策略和检索质量。"],
        ],
        [1.15, 1.5, 2.25, 1.6],
    )
    doc.add_heading("8.1 当前压测结果", level=2)
    add_callout(
        doc,
        "本地压测",
        "python scripts/load_test/simulate_api_load.py 1000 输出：requests=1000，served=135，rejected=865，cache_hit_rate=0.356，p50=3.549ms，p95=7.772ms。这个压测不是模拟真实生产 QPS，而是展示限流、缓存和延迟统计链路。",
    )
    doc.add_heading("8.2 缓存 key 设计逻辑", level=2)
    add_bullets(
        doc,
        [
            "天气查询 key：region_id + time_window + variables + data_version。",
            "塔点静态特征 key：tower_id + buffer_m + sensor_version + date_range。",
            "模型预测 key：model_version + tower_id/line_id + forecast_time + horizon。",
            "检索 key：query_hash + corpus_version + chunk_strategy + top_k。",
            "Agent 报告 key：prediction_version + guideline_version + output_schema_version。",
        ],
    )
    doc.add_heading("8.3 降级策略", level=2)
    add_matrix(
        doc,
        ["故障", "一级兜底", "二级兜底", "用户可见响应"],
        [
            ["CDS/GEE 临时失败", "使用最近成功下载数据", "切到 synthetic/demo 或规则模型", "提示数据时间戳和可信度下降。"],
            ["模型服务超时", "读取缓存预测", "运行轻量规则引擎", "返回风险等级但标注非最新模型输出。"],
            ["LLM/Agent 超时", "返回结构化规则解释", "只返回 CSV/HTML 基础报告", "避免生成无证据自然语言。"],
            ["检索无结果", "扩大 top_k 或换 chunk 策略", "返回标准应急模板", "提示未检索到匹配规程。"],
            ["高并发", "令牌桶限流", "按线路/区域批处理", "低优先级请求排队或重试。"],
        ],
        [1.25, 1.7, 1.65, 1.9],
    )
    doc.add_heading("8.4 RAG 切片策略差异", level=2)
    add_matrix(
        doc,
        ["切片策略", "适合文档", "优点", "风险"],
        [
            ["固定窗口 + overlap", "日志、长文本、没有清晰结构的资料", "实现简单，召回稳定。", "可能切断语义；chunk 太多导致成本高。"],
            ["段落/章节切片", "规程、手册、技术文档", "语义完整，方便引用。", "段落长度不均，可能召回不均衡。"],
            ["表格行/实体切片", "杆塔台账、巡检记录、设备缺陷表", "实体粒度清晰，适合精确查询。", "缺少上下文时解释不足。"],
            ["时间窗口切片", "告警流、天气序列、日志", "适合时序诊断。", "窗口边界会影响事件完整性。"],
            ["混合切片", "复杂企业知识库", "兼顾召回和语义。", "需要路由和评估，工程复杂。"],
        ],
        [1.25, 1.45, 1.8, 2.0],
    )
    doc.add_heading("8.5 Agent 多轮记忆取舍", level=2)
    add_bullets(
        doc,
        [
            "不是什么都记：用户身份、研究区、约束条件、已确认方案可以记；临时中间日志、重复输出不应长期记。",
            "短期记忆保留最近 turns，适合指代消解和连续任务。",
            "长期记忆做摘要，保留事实和决策，不保留冗余原文。",
            "证据不塞进上下文：原始 CSV、规程、历史报告应通过检索按需取回。",
            "上下文溢出时先压缩历史，再裁剪低优先级信息，最后要求用户确认关键缺失。",
        ],
    )


def add_interview_questions(doc: Document) -> None:
    doc.add_heading("九、岗位视角面试问题库", level=1)
    sections = {
        "AI Lab / AI for Science 技术面": [
            ("这个项目的科学问题是什么？", "核心是复杂山地中从区域格点气象到杆塔级灾害风险的尺度下推，以及多源地球系统数据与物理先验融合。"),
            ("为什么要加入 DLR？", "因为高影响天气预警不能停在 hazard 层。电网真正关心的是线路承载能力、运行约束和处置优先级，DLR 把气象变量映射到电网运行能力。"),
            ("为什么不用直接训练 GraphCast/Pangu 这类天气大模型？", "资源和目标都不合适。我的问题不是全球预报，而是行业风险下游适配。更合理的是使用 ERA5-Land/预训练天气模型输出作为上游气象场，再做区域微气象和线路风险建模。"),
            ("弱标签是否科学？", "弱标签不是最终真值，而是领域先验的可复现表达。它用于构建 MVP，后续可用灾害公报、巡检记录或仿真数据校准。面试中要主动承认其局限。"),
            ("如何提高科学可信度？", "做消融、案例验证、不确定性评估、专家规则对齐，并引入真实故障/气象灾害记录作为独立验证。"),
        ],
        "电网 / 能源算法面": [
            ("为什么选贵州？", "贵州山地高海拔、冬季低温高湿、凝冻灾害有现实意义，且复杂地形能突出 DEM 和下垫面特征的价值。"),
            ("真实电网数据涉密怎么办？", "公开项目使用合成线路或公开 OSM 线路，模型链路和特征工程可复现；如果进入企业环境，可替换为真实杆塔和故障标签。"),
            ("业务上最重要的指标是什么？", "高风险召回、提前预警时间、误报/漏报成本、线路段级连续风险，而不只是总体 accuracy。"),
            ("风险和动态载流量是什么关系？", "风险描述外部灾害概率，DLR 描述线路热容量余量。二者结合才能区分“天气危险但运行影响小”和“天气危险且容量受限”的场景。"),
            ("怎么解释某个杆塔高风险？", "展示温度、湿度、降水、风速、海拔、坡度、NDVI 等触发因子，输出可追溯解释。"),
        ],
        "大厂算法岗": [
            ("为什么先用 GBDT 类模型？", "表格特征强 baseline，训练快、稳定、可解释。先证明数据和标签链路有效，再上深度时序模型。"),
            ("如何避免数据泄漏？", "按时间切分训练/测试，预处理封装在 pipeline 中，静态特征不使用未来信息，GEE 指数用历史窗口聚合。"),
            ("如何做消融实验？", "weather only、weather + DEM、weather + DEM + Sentinel、+ line geometry、+ physics labels，对比 macro-F1 和 high-risk recall。"),
            ("下一步模型升级？", "LightGBM/XGBoost、TFT/PatchTST、图神经网络线路拓扑、概率预测和校准。"),
        ],
        "Agent 岗": [
            ("这个 Agent 和普通 Chatbot 有什么区别？", "它是 tool-using Agent，调用风险模型、气象查询、塔点查询和规则解释，不凭语言模型直接生成结论。"),
            ("如何减少幻觉？", "所有解释来自结构化变量和规则；报告中显示证据字段；无法获取数据时拒答或提示缺失。"),
            ("RAG 可以加在哪里？", "加入电网运维规程、覆冰应急预案、历史案例库，检索后生成处置建议。"),
            ("如何评估 Agent？", "工具调用成功率、证据一致性、建议可执行性、人工审核通过率。"),
        ],
        "多模态岗": [
            ("多模态体现在哪里？", "小时级气象序列、空间遥感/DEM raster、线路点线结构是三类不同模态。"),
            ("为什么不是简单拼接？", "不同模态时间尺度不同：气象是动态，DEM/Sentinel 是静态，线路是结构属性。融合时要注意时空对齐和变量语义。"),
            ("遥感预训练模型怎么加？", "用 DOFA/Prithvi/RS-CLIP 等提取 tile embedding，与气象时序 encoder 融合。"),
            ("如何证明遥感有用？", "做消融，看加入 NDVI/NDWI/NDBI/DEM 后 high-risk recall 是否提升。"),
        ],
        "AI Infra / MLOps 岗": [
            ("这个项目如何服务化？", "FastAPI 暴露预测接口，PostGIS 存塔点和线路，Redis 缓存 GEE/气象特征，Docker 部署，定时任务下载新气象数据。"),
            ("如何处理大范围批量推理？", "按空间 tile 和时间窗口分批，异步队列，模型常驻内存，结果写入 parquet/数据库。"),
            ("如何监控？", "记录下载成功率、数据延迟、推理耗时、异常比例、高风险数量漂移、模型输入分布漂移。"),
            ("为什么要把数据、模型、报告分层？", "便于复现、调试、缓存和替换模型；真实工程里最怕一步脚本黑箱到底。"),
        ],
        "HR / 综合面": [
            ("你为什么做这个项目？", "我希望把自己遥感和地球系统数据背景拓展到能源行业落地，证明自己既能做科学问题，也能做工程闭环。"),
            ("项目最大的困难是什么？", "真实困难不是模型，而是数据权限、CDS/GEE 认证、代理、格式兼容、非敏感线路构造和多源对齐。"),
            ("你在团队中能承担什么角色？", "我能做数据工程、算法建模、领域分析和结果解释之间的桥接角色。"),
            ("项目有什么不足？", "当前没有真实电网故障标签，弱标签需要进一步用企业数据或公开灾害案例校准；深度时序模型和服务化部署还可增强。"),
        ],
    }
    for title, qas in sections.items():
        doc.add_heading(title, level=2)
        for q, a in qas:
            add_para(doc, f"Q：{q}")
            add_para(doc, f"A：{a}")


def add_math_cs_questions(doc: Document) -> None:
    doc.add_heading("十、基础知识高频追问", level=1)
    add_matrix(
        doc,
        ["类别", "问题", "回答要点"],
        [
            ["计算机", "为什么用配置文件而不是硬编码？", "实验可复现；研究区、时间、变量、路径可替换；便于部署和 CI。"],
            ["计算机", "如何保证 pipeline 可复现？", "固定随机种子、保存配置、数据 schema 稳定、训练/推理同一预处理 pipeline。"],
            ["计算机", "TTL-LRU 缓存复杂度？", "哈希表 + 双向链表/OrderedDict，get/set 平均 O(1)，超时惰性删除。"],
            ["计算机", "令牌桶如何支持突发流量？", "桶容量允许短时 burst，refill rate 控制长期平均速率。"],
            ["计算机", "限流和降级区别？", "限流控制进入系统的请求，降级是在资源不足或依赖失败时返回低成本可接受结果。"],
            ["计算机", "RAG 检索为什么要压测？", "检索质量和延迟都会影响 Agent，必须看 recall@k、空结果率、p95、rerank 成本。"],
            ["数学", "为什么 macro-F1 比 accuracy 更重要？", "风险等级不均衡，高风险小类容易被 accuracy 掩盖。macro-F1 对每类等权。"],
            ["数学", "GBDT 的基本思想？", "逐步拟合损失函数负梯度，用多个弱学习器加和形成强模型。"],
            ["数学", "时间切分为什么重要？", "随机切分会把相邻时间的天气状态同时放入训练和测试，造成时间泄漏。"],
            ["遥感", "NDVI/NDWI/NDBI 分别是什么？", "植被、水体/湿润、建筑/裸地 proxy，来自不同光谱波段比值。"],
            ["遥感", "为什么用 buffer 均值？", "杆塔风险受周边下垫面影响，单点像元容易受噪声和配准误差影响。"],
            ["气象", "ERA5-Land 和观测站有什么区别？", "ERA5-Land 是再分析格点产品；观测站是真实点观测但空间稀疏。"],
            ["气象", "相对湿度如何从温度和露点估算？", "用饱和水汽压公式 e/es，露点越接近温度，相对湿度越高。"],
            ["电网", "覆冰风险为什么受风影响？", "风影响水滴输送、碰撞效率和导线迎风侧沉积。"],
            ["电网", "动态线路评级的物理依据？", "导线热平衡：焦耳热和太阳辐射输入，与对流/辐射散热平衡；风和低温通常提高允许电流。"],
        ],
        [0.75, 1.75, 4.0],
    )


def add_runbook_and_code_walkthrough(doc: Document) -> None:
    doc.add_heading("十一、可复现实验 Runbook 与代码级 Walkthrough", level=1)
    doc.add_heading("9.1 一键复现实验命令", level=2)
    add_matrix(
        doc,
        ["目标", "命令", "输出", "面试说明"],
        [
            ["Demo 全流程", "python scripts/run_demo_pipeline.py", "demo training table、模型、预测、HTML 报告", "没有账号也能复现项目闭环，适合 GitHub 展示。"],
            ["下载 ERA5-Land", "python scripts/download_era5_land.py --config configs/project.yaml", "data/demo/raw/era5_land/reanalysis-era5-land_2023_01.nc", "CDS 任务会排队，脚本按月份请求。"],
            ["转换 ERA5", "python scripts/prepare_era5_weather.py --input data/demo/raw/era5_land --output data/real/raw/weather_hourly.csv", "标准 weather_hourly.csv", "兼容 CDS 返回 ZIP-in-NC 的情况。"],
            ["提取 GEE 静态特征", "python scripts/export_gee_static_features.py --auth-mode user --ee-project nmcproductivity --proxy-url http://127.0.0.1:7897 ...", "tower_static_features_gee.csv", "GEE 认证、project、proxy 是真实工程问题。"],
            ["真实数据训练", "python scripts/run_real_era5_pipeline.py", "真实 ERA5 + GEE 特征训练指标和报告", "当前最终主链路。"],
            ["单元测试", "python -m pytest tests", "1 passed", "测试物理标签逻辑基本正确性。"],
            ["压测", "python scripts/load_test/simulate_api_load.py 1000", "cache hit rate、p50/p95、rejected", "展示上线意识和限流缓存设计。"],
            ["消融", "python scripts/experiments/run_ablation.py", "ablation_results.csv/json", "回答为什么加 DEM/Sentinel/DLR/IEEE738。"],
            ["解释图", "python scripts/experiments/permutation_importance.py", "permutation_importance.png/csv", "回答模型怎么看特征贡献。"],
            ["检索评估", "python scripts/retrieval_eval/evaluate_chunking.py", "chunking_eval.json", "回答 chunking 和 hybrid retrieval。"],
            ["PatchTST-lite", "python scripts/experiments/train_patchtst_lite.py", "patchtst_lite_metrics.json", "环境不支持 torch 时写 skipped，不伪造结果。"],
        ],
        [1.0, 2.25, 1.55, 1.7],
    )
    doc.add_heading("9.2 关键代码文件逐行讲法", level=2)
    add_matrix(
        doc,
        ["文件", "关键函数/类", "做了什么", "容易被问到"],
        [
            ["config.py", "load_config / project_path / ensure_dirs", "统一读取 YAML 配置，避免路径和参数散落在脚本里。", "为什么配置化？如何管理多实验？"],
            ["era5_downloader.py", "Era5Request / monthly_requests / download_era5_land", "将时间窗拆成月度请求，调用 cdsapi.retrieve 下载。", "CDS 任务失败如何重试？如何避免重复下载？"],
            ["era5_to_weather.py", "convert_era5_folder", "自动识别 ZIP/NetCDF，读取变量并转换单位。", "ERA5 变量维度、单位、经纬度命名差异。"],
            ["export_gee_static_features.py", "export_static_features", "GEE 初始化、构造 tower buffer、计算 DEM 和 Sentinel 指数、reduceRegions。", "GEE 权限、代理、云端计算和本地下载区别。"],
            ["synthetic.py", "generate_synthetic_towers", "生成非敏感线路点，保留线路方向和空间分布。", "为什么不直接用真实电网数据？"],
            ["build_dataset.py", "_nearest_weather_for_towers / build_training_table", "将杆塔匹配到最近气象格点，并合并动态和静态特征。", "最近邻匹配误差如何改进？"],
            ["dlr.py", "add_dynamic_line_rating_features", "用温度、风速、线路夹角估算动态线路容量和热压力。", "DLR 代理和正式标准模型差距在哪里？"],
            ["model_zoo.py", "make_classifier", "LightGBM/XGBoost 可选模型工厂，未安装时 fallback。", "如何设计可选依赖？"],
            ["graph_smooth_predictions.py", "smooth_line_graph", "同一线路相邻杆塔风险平滑，模拟图拓扑消息。", "GNN 为什么适合线路拓扑？"],
            ["infra/cache.py", "TTLLRUCache", "TTL + LRU 缓存，控制延迟和内存。", "手撕 LRU/TTL cache。"],
            ["infra/rate_limit.py", "TokenBucketRateLimiter", "令牌桶限流，支持突发流量。", "手撕限流器。"],
            ["infra/circuit_breaker.py", "CircuitBreaker", "连续失败后打开熔断，恢复后半开探测。", "熔断器状态机。"],
            ["infra/singleflight.py", "SingleFlight", "防止同 key 并发穿透，减少缓存击穿。", "如何避免多个线程同时打下游？"],
            ["retrieval/chunking.py", "fixed_window_chunks / paragraph_chunks", "比较不同切片策略。", "chunk size/overlap 如何选？"],
            ["retrieval/bm25.py", "BM25Retriever", "轻量关键词检索 baseline。", "BM25 和向量检索区别？"],
            ["retrieval/hybrid.py", "HybridRetriever / simple_rerank", "BM25 + TF-IDF 向量混合检索，再用覆盖率 rerank。", "hybrid search 为什么比单一路线稳？"],
            ["agent/memory.py", "BudgetedConversationMemory", "摘要 + 最近 turns 控制上下文预算。", "如何避免上下文溢出？"],
            ["service/api.py", "FastAPI app", "健康检查、单塔预测、缓存和限流。", "接口延迟和部署依赖如何处理？"],
            ["physics_labels.py", "compute_icing_risk_score / add_physics_labels", "把气象灾害先验转成可复现弱标签。", "弱标签偏差、阈值选择、物理合理性。"],
            ["train.py", "train_risk_model", "时间切分、预处理 pipeline、分类器和回归器训练。", "时间泄漏、类别不均衡、为什么 macro-F1。"],
            ["predict.py", "predict_latest", "取最新 24h 样本批量预测风险等级和分数。", "线上推理如何做增量更新？"],
            ["explain.py/report.py", "explain_row / build_html_report", "将风险触发因子写成可读解释并生成报告。", "规则解释和 SHAP、LLM Agent 的关系。"],
        ],
        [1.25, 1.35, 2.05, 1.85],
    )
    doc.add_heading("9.3 当前生成文件解释", level=2)
    add_bullets(
        doc,
        [
            "data/real/raw/weather_hourly.csv：真实 ERA5-Land 标准气象表，包含 time、lat、lon、temperature_c、relative_humidity、wind_speed_ms、wind_dir_deg、precip_mm、pressure_hpa。",
            "data/real/raw/towers.csv：贵州研究区内的非敏感杆塔点，包含 tower_id、line_id、lat、lon、line_heading_deg。",
            "data/real/features/tower_static_features_gee.csv：GEE 提取的真实 DEM/Sentinel 静态特征，包含 elevation_m、slope_deg、aspect_deg、NDVI、NDWI、NDBI。",
            "data/real/features/training_table.csv：模型训练主表，每一行是某个时刻某个杆塔的气象、地形、遥感、线路和标签信息。",
            "artifacts/models_real_era5/risk_model.joblib：训练好的分类/回归模型 bundle，包含预处理 pipeline、模型和 metrics。",
            "artifacts/predictions_real_era5/latest_predictions.csv：最新 24 小时所有杆塔预测结果。",
            "artifacts/reports/gridweather_real_era5_report.html：面试展示用风险报告，可展示高风险杆塔和解释。",
        ],
    )


def add_deep_dive_questions(doc: Document) -> None:
    doc.add_heading("十二、深挖追问：技术面高压问题与回答框架", level=1)
    deep_qas = [
        ("如果面试官说你的标签是规则生成的，不是真实标签，这个项目还有价值吗？", "有价值，但要准确定位。当前版本是一个真实数据工程和弱监督原型，解决了公开环境下电网涉密标签缺失的问题。它证明了数据链路、特征融合、物理先验和报告闭环。生产版本需要接入企业脱敏故障标签或灾害案例做监督校准。回答时不要把弱标签包装成真实故障标签。"),
        ("为什么 ERA5-Land 分辨率仍然不够杆塔级？", "ERA5-Land 约 0.1 度，仍是公里级格点。杆塔级风险需要地形、坡向、海拔和线路方向做下推。当前项目的思想是：上游格点气象提供大尺度背景，DEM/Sentinel/线路特征提供局地修正。"),
        ("最近邻匹配天气格点有什么问题？", "最近邻简单稳定，但会造成空间离散误差。后续可改为双线性插值、海拔 lapse-rate 修正、风场地形修正，或者训练一个 downscaling 模型。"),
        ("为什么高风险 3 类召回低于其他类？", "最高风险样本通常更少、边界更窄，且弱标签中极端风险受多因子组合影响。后续可通过 class_weight、重采样、阈值移动、focal loss 或把高风险合并做二阶段模型提升召回。"),
        ("你如何做不确定性？", "可以做三层：输入不确定性，用多源预报/集合预报；模型不确定性，用 bootstrap/ensemble；阈值不确定性，用概率校准和 conformal prediction。业务输出应给风险等级 + 置信区间。"),
        ("如果上线，每小时新数据来了怎么增量更新？", "把静态 GEE 特征缓存到数据库，动态气象按时间追加，最近 24-72h 构建推理 batch，模型服务只读取最新窗口，不重复计算静态特征。"),
        ("Agent 部分是否太简单？", "当前是规则解释 Agent 的 MVP。它的价值是可靠和可追溯。升级方向是工具调用 Agent：query_weather、query_tower、run_model、retrieve_guideline、generate_report，每一步都有结构化证据。"),
        ("DLR 模块是否过于简化？", "是简化版，所以文档里明确叫 DLR proxy，不冒充认证级工程计算。它的价值是把天气风险和电网运行能力连接起来。后续可以实现 IEEE 738/CIGRE 热平衡，加入太阳辐射、导线型号、实测导线温度和 SCADA 负荷。"),
        ("为什么不用 LLM 直接判断风险？", "LLM 不适合直接做数值风险预测。风险预测应由可验证模型和物理规则完成，LLM/Agent 负责调用工具、组织证据和生成报告。"),
        ("怎么证明多模态融合有效？", "必须做消融实验，而不是口头说有效。至少比较 weather only、weather+DEM、weather+DEM+Sentinel、weather+DEM+Sentinel+line geometry。"),
        ("如果真实电网给了故障标签，你怎么改？", "将 weak label 变为 auxiliary target 或 regularization，主目标改为真实故障/覆冰厚度；做时间/空间外推验证；关注 high-risk recall、PR-AUC、提前量。"),
        ("如果接口 QPS 突然升高怎么办？", "先限流保护核心服务；高频 tower/time 查询走缓存；低优先级报告生成排队；LLM 解释降级为规则解释；重要线路优先处理。"),
        ("如何控制 LLM token 成本？", "只让 LLM 生成解释和报告，不让它处理全量数值；检索 top-k 控制；长历史摘要化；同类报告模板化；缓存相同 query 和相同预测版本的输出。"),
        ("RAG chunking 怎么选？", "运维规程按章节/段落，台账按实体行，告警日志按时间窗口，长文本用固定窗口+overlap。最终用评估集比较 recall@k、MRR 和延迟。"),
        ("Agent 记忆怎么避免污染？", "只写入稳定事实和用户确认决策；工具输出带版本和时间戳；错误中间结果不进入长期记忆；支持记忆过期和人工清除。"),
        ("如何做检索压测？", "构造典型问题集和标准答案片段，测 recall@k、MRR、p95 latency、空结果率、rerank 成本；对不同 chunk size、overlap、top_k 做网格实验。"),
    ]
    for q, a in deep_qas:
        add_callout(doc, f"Q：{q}", f"A：{a}")


def add_hr_pack(doc: Document) -> None:
    doc.add_heading("十三、HR 面与行为面回答素材", level=1)
    add_matrix(
        doc,
        ["问题", "回答结构", "可直接使用的表达"],
        [
            ["为什么做这个项目？", "背景-目标-行动-结果", "我原本有遥感和地球系统数据背景，但秋招目标不只限科研岗，所以做了一个能源行业落地项目，证明自己能把真实数据、模型和报告系统打通。"],
            ["你最大的优势？", "交叉能力", "我不是只会调模型，也不是只懂领域。我能把遥感、气象、地理空间数据、物理机制和工程实现串起来。"],
            ["遇到最大困难？", "真实工程问题", "不是模型本身，而是 CDS/GEE 权限、代理、格式兼容、真实标签不可公开和多源数据对齐。这些问题解决后，工程才真正可复现。"],
            ["项目中你的贡献？", "端到端 owner", "我从问题定义、数据源选择、下载脚本、GEE 特征、弱标签、模型、报告到面试包装都独立完成。"],
            ["项目不足？", "诚实承认 + 计划", "当前还没有真实电网故障标签，弱标签需要企业数据或公开案例校准；后续会补消融、SHAP、服务化和真实案例验证。"],
            ["为什么适合我们岗位？", "岗位映射", "这个项目覆盖数据工程、算法建模、业务解释和系统化交付，和 AI Lab/电网/大厂算法/Agent/AI Infra 的交叉要求一致。"],
        ],
        [1.25, 1.25, 4.0],
    )
    doc.add_heading("STAR 案例：解决 GEE 权限和代理问题", level=2)
    add_para(
        doc,
        "Situation：项目需要从 GEE 提取 Sentinel-2/NASADEM 静态特征，但本地一开始出现 SSL EOF、project 未注册、service account 无权限、OAuth token 过期等问题。Task：保证 GEE 自动提取能够接入主工程。Action：逐步定位为代理端口、Earth Engine project、service account IAM 和用户认证问题；最终使用 nmcproductivity 注册项目、刷新用户认证，并为脚本增加 --ee-project、--proxy-url、auth-mode user/service-account 和 requests proxy patch。Result：成功生成 140 个贵州塔点的真实 GEE 静态特征，并接入真实 ERA5 pipeline，macro-F1 提升到 0.877。",
    )
    doc.add_heading("STAR 案例：从概念方案到可运行工程", level=2)
    add_para(
        doc,
        "Situation：原始方案包含 Pangu-Lite、RS-CLIP、Cross-Attention、FAT、Slurm 等很多概念，但风险是不可落地。Task：在本机条件下做出秋招可展示项目。Action：收缩问题到贵州山地输电线路微气象风险，优先完成公开数据、非敏感线路、弱标签、强 baseline 和报告闭环。Result：项目从概念变为可运行工程，真实数据版本包含 70,560 条训练样本和端到端报告。",
    )


def add_resume_and_roadmap(doc: Document) -> None:
    doc.add_heading("十四、简历写法与后续迭代路线", level=1)
    doc.add_heading("14.1 简历项目描述", level=2)
    add_callout(
        doc,
        "简历版",
        "基于 Copernicus ERA5-Land、Google Earth Engine Sentinel-2/NASADEM 和非敏感合成输电线路，构建贵州山地 Weather-to-Grid Resilience Agent，实现数据下载、遥感地形特征提取、物理弱标签生成、动态线路容量余量估计、风险模型训练、杆塔级推理和可解释动作建议报告；补充 TTL-LRU 缓存、令牌桶限流、异常降级、RAG 切片检索、Agent 记忆预算和本地压测模块，体现从算法 demo 到可上线工程的完整设计。",
    )
    doc.add_heading("14.2 后续增强路线", level=2)
    add_numbers(
        doc,
        [
            "加入公开灾害案例或企业脱敏故障标签，替换/校准弱标签。",
            "系统消融已实现；后续可把 ablation 扩展到多月份、多区域和真实灾害案例。",
            "Permutation importance 已实现；后续可安装 SHAP 做局部解释和 force plot。",
            "PatchTST-lite 接口已实现；后续修复 torch 环境或迁移到 Linux/GPU 后训练深度时序模型。",
            "IEEE738-like DLR 已实现；后续加入导线型号、太阳辐射实测、导线温度、SCADA 负荷，升级为认证级热平衡。",
            "FastAPI/Docker/PostGIS/Redis 配置已加入；后续安装依赖后做真正服务启动和压测。",
            "Agent/RAG 已加入 demo corpus 和 hybrid retrieval；后续接入真实运维规程、reranker 和答案评估集。",
            "压测已具备缓存、限流、熔断、防击穿；后续加入并发线程、队列削峰和灰度发布。",
            "检索增强已加入 hybrid search 和 chunking eval；后续做向量 embedding、reranker 和自动 chunk 网格搜索。",
        ],
    )
    doc.add_heading("14.3 面试自我介绍 60 秒版本", level=2)
    add_para(
        doc,
        "我最近做了一个 GridWeather-Agent 项目，目标是把遥感、气象和地理空间数据落到电网极端天气韧性运行场景。项目选取贵州毕节-六盘水山地走廊，自动下载 ERA5-Land 气象数据，并通过 Google Earth Engine 提取 NASADEM 和 Sentinel-2 静态特征；考虑真实电网数据涉密，我用非敏感合成线路模拟杆塔分布，再用低温、高湿、降水、风速和地形暴露构造物理弱标签，同时加入动态线路容量余量代理，把天气风险连接到运行约束。当前真实数据版本有 70,560 条训练样本，macro-F1 达到 0.879，并能输出每个杆塔的风险解释、DLR 余量和动作建议。这个项目体现了我从科学问题、数据工程、模型训练到可解释系统落地的完整能力。",
    )


def add_coding_interview_section(doc: Document) -> None:
    doc.add_heading("十五、项目相关手撕代码题", level=1)
    add_para(doc, "下面这些题和本项目强相关，适合算法面、后端工程面、AI Infra 面。面试时先说思路、复杂度，再写代码。")
    problems = [
        (
            "手撕 TTL + LRU Cache",
            "题意：实现 get/set，容量满时淘汰最久未使用 key，过期 key 返回 None。",
            "思路：OrderedDict 或 hash map + doubly linked list。get 时检查 TTL，命中后 move_to_end；set 时更新过期时间，超过容量 pop oldest。复杂度 O(1)。",
            "项目对应：gridweather.infra.cache.TTLLRUCache。",
        ),
        (
            "手撕令牌桶限流器",
            "题意：capacity 表示桶容量，rate 表示每秒补充 token，allow(cost) 判断是否放行。",
            "思路：每次请求先按 elapsed_time * rate 补充 token，再判断 tokens >= cost。支持突发流量，长期平均受 rate 控制。",
            "项目对应：gridweather.infra.rate_limit.TokenBucketRateLimiter。",
        ),
        (
            "手撕固定窗口 chunking",
            "题意：给长文档和 size/overlap，输出 chunk 列表。",
            "思路：step = size - overlap，从 0 开始滑窗，注意 overlap < size，否则死循环。复杂度 O(n)。",
            "项目对应：gridweather.retrieval.chunking.fixed_window_chunks。",
        ),
        (
            "手撕 BM25 检索",
            "题意：给文档集合和 query，返回 top-k 文档。",
            "思路：统计 tf、df、avgdl，score = idf * tf*(k1+1)/(tf+k1*(1-b+b*dl/avgdl))。适合关键词召回。",
            "项目对应：gridweather.retrieval.bm25.BM25Retriever。",
        ),
        (
            "手撕滑动窗口最大值/最近 24h 聚合",
            "题意：对每个杆塔时序计算最近窗口最大风险或累计降水。",
            "思路：按 tower_id 分组，按时间排序；固定窗口可用 deque 维护最大值，复杂度 O(n)。",
            "项目延伸：预测最新 24h 风险和告警聚合。",
        ),
        (
            "手撕 top-k 高风险杆塔",
            "题意：从海量预测结果中找风险分数最高的 k 个杆塔。",
            "思路：小顶堆维护 k 个元素，复杂度 O(n log k)，比全排序 O(n log n) 更适合流式场景。",
            "项目延伸：报告里 top risk towers。",
        ),
        (
            "手撕重试与退避",
            "题意：调用外部 API 失败时重试，间隔指数增长。",
            "思路：捕获可重试异常，sleep(base * 2^i)，超过次数抛出或 fallback。",
            "项目对应：gridweather.infra.resilience.retry。",
        ),
        (
            "手撕 Agent 记忆裁剪",
            "题意：给多轮对话和 token 预算，保留最近 turns，旧信息摘要化。",
            "思路：recent window + summary；超过预算时先摘要最旧消息，再截断摘要。项目中估算 token。",
            "项目对应：gridweather.agent.memory.BudgetedConversationMemory。",
        ),
    ]
    for title, desc, idea, mapping in problems:
        add_callout(doc, title, f"{desc}\n\n{idea}\n\n{mapping}")


def build() -> Path:
    doc = Document()
    set_styles(doc)
    add_title_page(doc)
    add_toc(doc)
    add_industry_need(doc)
    add_project_story(doc)
    add_architecture(doc)
    add_data_steps(doc)
    add_algorithms(doc)
    add_domain_knowledge(doc)
    add_infra_agent(doc)
    add_production_engineering(doc)
    add_interview_questions(doc)
    add_math_cs_questions(doc)
    add_runbook_and_code_walkthrough(doc)
    add_deep_dive_questions(doc)
    add_hr_pack(doc)
    add_resume_and_roadmap(doc)
    add_coding_interview_section(doc)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(OUT)
        return OUT
    except PermissionError:
        fallback = OUT.with_name("GridWeather-Agent_秋招面试项目详解手册_v2.docx")
        doc.save(fallback)
        return fallback


if __name__ == "__main__":
    print(build())
